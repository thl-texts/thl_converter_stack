#!env/bin/python
"""
A convert that just addes digital pages to word docs
"""
from baseconverter import BaseConverter
from os import path
import docx
from docx.text.run import Run
from docx.oxml.text.run import CT_R
import re
import logging


class DigitalPages(BaseConverter):
    outfile = ""
    lines_per_page = 15
    tsek_per_line = 20
    digstyles = {
        'page': 'page number', # This is the digital page style name (page number repurposed)
        'line': 'line number' # this is the digital line style name
    }
    # style_id_list = [digpage_style_id, digline_style_id]
    pgct = 1  # Page start 1 (to insert first page and line markers)
    lnct = 0  # Line start 0 (increment after adding so we can use = below)
    ct = 0
    tmplts = {
        'page': 'tdp',
        'line': 'tdl'
    }
    tsekpattern = r'[\u0F00-\u0F14\u0F3A-\u0F3D\u0FD2-\u0FD8\s]+'  # at least one of the Tibetan punctuation or space-like characters
    tskcount = 0
    donefirst = False

    def convertdoc(self):
        self.current_file_path = path.join(self.indir, self.current_file)
        self.worddoc = docx.Document(self.current_file_path)
        self.merge_runs()
        totalpct = len(self.worddoc.paragraphs)
        pct = 0
        # Reset class variables for each document
        self.pgct = 1
        self.lnct = 0
        self.ct = 0
        self.tskcount = 0
        self.donefirst = False

        print("Inserting placeholders ...")

        foundhead = False
        for p in self.worddoc.paragraphs:
            pct += 1
            print("\rDoing paragraph {} of {}        ".format(pct, totalpct), end="")
            psnm = p.style.name
            # Don't start counting until we get the first header (usually front or body)
            if not foundhead and 'Heading' not in psnm:
                continue
            foundhead = True
            # After that ignore headers and only proc ess non-headers
            if 'Heading' not in psnm:
                self.insert_ms(p)
                self.apply_styles(p)
        print("\n")
        self.outfile = path.join(self.outdir, self.current_file.replace('.doc', '-out.doc'))
        self.worddoc.save(self.outfile)

    def insert_ms(self, p):
        for r in p.runs:
            inserts = []
            rtxt = r.text
            endindex = -1
            for m in re.finditer(self.tsekpattern, rtxt):
                endindex = m.end()
                if m.start() == 0:
                    continue
                self.tskcount += 1
                if self.tskcount == self.tsek_per_line:
                    self.lnct += 1
                    self.tskcount = 0
                    if self.lnct == self.lines_per_page:
                        self.pgct += 1
                        self.lnct = 0
                        inserts.append(('page', f"{self.pgct}", m))
                        inserts.append(('line', f"{self.pgct}.{self.lnct + 1}", m))
                    else:
                        inserts.append(('line', f"{self.pgct}.{self.lnct + 1}", m))

            inserts.reverse()
            for ins in inserts:
                mstype, msnum, m = ins
                ms = f"[{self.tmplts[mstype]} {msnum}]"
                rtxt = rtxt[:m.end()] + ms + rtxt[m.end():]

            r.text = rtxt
            if not self.donefirst:
                rtxt = f"[{self.tmplts['page']} 1][{self.tmplts['line']} 1.1]" + rtxt
                r.text = rtxt
                self.donefirst = True

    def apply_styles(self, p):
        newp = p.insert_paragraph_before('', p.style)
        for r in p.runs:
            rtxt = r.text
            pgrex = '\[' + self.tmplts["page"] + '\s+\d+\]'
            mtchs = [m for m in re.finditer(pgrex, rtxt)]
            if len(mtchs) > 0:
                newp.add_run(rtxt[:mtchs[0].start()], r.style)
                for mi, mtch in enumerate(mtchs):
                    ms = rtxt[mtch.start():mtch.end()]
                    ms = ms.replace(self.tmplts["page"] + ' ', '')
                    newp.add_run(ms, self.worddoc.styles[self.digstyles['page']])
                    endindex = mtchs[mi + 1].start() if mi < len(mtchs) - 1 else len(rtxt)
                    postrun = rtxt[mtch.end():endindex]
                    newp.add_run(postrun, r.style)
            else:
                newp.add_run(r.text, r.style)
        self.delete_paragraph(p)
        # Do Line milestones
        newp2 = newp.insert_paragraph_before('', newp.style)
        for r in newp.runs:
            rtxt = r.text
            pgrex = '\[' + self.tmplts["line"] + '\s+\d+\.\d+\]'
            mtchs = [m for m in re.finditer(pgrex, rtxt)]
            if len(mtchs) > 0:
                # Add text before first match
                newp2.add_run(rtxt[:mtchs[0].start()], r.style)
                for mi, mtch in enumerate(mtchs):
                    # add milestone run
                    ms = rtxt[mtch.start():mtch.end()]
                    ms = ms.replace(self.tmplts["line"] + ' ', '')
                    newp2.add_run(ms, self.worddoc.styles[self.digstyles['line']])
                    # Add post milestone run
                    endindex = mtchs[mi + 1].start() if mi < len(mtchs) - 1 else len(rtxt)
                    postrun = rtxt[mtch.end():endindex]
                    newp2.add_run(postrun, r.style)
            else:
                newp2.add_run(r.text, r.style)
        self.delete_paragraph(newp)

    def add_run(self, r, txt, style_name):
        runel = CT_R()
        r._element.addnext(runel)
        newrun = Run(runel, r._parent)
        newrun.text = txt
        newrun.style = self.worddoc.styles[style_name]
        return newrun

    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None


####### OLD CODE ########
class OldDigitalPages(BaseConverter):
    outfile = ""
    lines_per_page = 15
    tsek_per_line = 20
    digpage_sigla = 'TDP'  # for Tibetan Digital Page
    digline_sigla = 'TDL'  # for Tibetan Digital Line
    digpage_style_id = 'page number'
    digline_style_id = 'line number'
    style_id_list = [digpage_style_id, digline_style_id]

    def convertdoc(self):
        self.current_file_path = path.join(self.indir, self.current_file)
        self.worddoc = docx.Document(self.current_file_path)

        # stylist = []
        # for st in self.worddoc.styles:
        #     stylist.append("{} | {}".format(st.name, st.style_id))
        #     if st.name.lower() == 'line number' or st.name.lower() == 'LineNumber':
        #         print(st.name)
        # stylist.sort()
        # for st in stylist:
        #     print(st)
        # exit(0)

        pgct = 0                    # Page start 0
        lnct = 1                    # Line start 1
        ct = 0
        tsekpattern = r'[ཀ-࿿]་|[ཀ-࿿]།|ག\s'
        pct = 0
        totalpct = len(self.worddoc.paragraphs)
        print("Inserting milestones ...")
        for p in self.worddoc.paragraphs:
            pct += 1
            print("\rDoing paragraph {} of {}        ".format(pct, totalpct), end="")
            psnm = p.style.name
            if 'Heading' not in psnm:
                # print(psnm, len(p.runs))
                for r in p.runs:
                    insertPoints = []
                    rstyname = r.style.style_id
                    if "Default" not in rstyname:
                        logging.log(logging.DEBUG, "Style ID: {}".format(rstyname))

                    if pgct == 0:
                        pgct += 1
                        insertPoints.append({
                            'page': pgct,
                            'line': lnct,
                            'match': 0
                        })
                    tseklist = re.finditer(tsekpattern, r.text)
                    for tsek in tseklist:
                        ct += 1
                        if ct == self.tsek_per_line:
                            if lnct == self.lines_per_page:
                                pgct += 1
                                lnct = 1
                                insertPoints.append({
                                    'page': pgct,
                                    'line': lnct,
                                    'match': tsek
                                })
                            else:
                                lnct += 1
                                insertPoints.append({
                                    'page': pgct,
                                    'line': lnct,
                                    'match': tsek
                                })
                            ct = 0
                    insertPoints.reverse()
                    rtxt = r.text
                    for ip in insertPoints:
                        stind = ip['match'] if isinstance(ip['match'], int) else ip['match'].span()[1]
                        lnnum = ip['line']
                        pgnum = ip['page']
                        ms = ''
                        if lnnum == 1:
                            ms += '[{} {}]'.format(self.digpage_sigla, pgnum)
                        ms += '[{} {}.{}]'.format(self.digline_sigla, pgnum, lnnum)
                        rtxt = rtxt[:stind] + ms + rtxt[stind:]
                    r.text = rtxt

        # Need to go through document again and find all [DP 1] and [DL 13] etc. and apply character styles:
        # "Page Number" for pages and "Line Number" for lines. (Check if these are the standard style names)
        # Style ID: PageNumber and  Style ID: LineNumber
        pi = 0
        digpage_style = self.worddoc.styles[self.digpage_style_id]
        digline_style = self.worddoc.styles[self.digline_style_id]

        digmile_match = r'\[({})\s+\d+\]|\[({})\s+\d+\.\d+\]'.format(self.digpage_sigla, self.digline_sigla)
        ct = 0
        totalct = len(self.worddoc.paragraphs)
        print("\nApplying styles ...")
        for p in self.worddoc.paragraphs:
            ct += 1
            print("\rparagraph {} of {}        ".format(ct, totalct), end="")
            psnm = p.style.name
            if 'Heading' not in psnm:
                myruns = p.runs  # get all runs
                p.clear()        # remove runs to reinsert with splits
                for r in myruns:
                    rtxt = r.text
                    rsty = r.style
                    if r.style.name not in self.style_id_list:
                        mtch = re.search(digmile_match, rtxt)
                        while mtch:
                            (mst, men) = mtch.span()
                            mtxt = mtch.group()
                            sig = mtch.group(1)
                            milesty = digpage_style if sig == self.digpage_sigla else digline_style
                            beftxt = rtxt[:mst]
                            afttxt = rtxt[men:]
                            if len(beftxt) > 0:
                                p.add_run(beftxt, rsty)
                            p.add_run(mtxt, milesty)
                            rtxt = afttxt
                            mtch = re.search(digmile_match, rtxt)
                        if len(rtxt) > 0:
                            p.add_run(rtxt, rsty)
                    else:
                        p.add_run(rtxt, rsty)

        print("\n")
        self.outfile = path.join(self.outdir, self.current_file.replace('.doc', '-out.doc'))
        self.worddoc.save(self.outfile)
