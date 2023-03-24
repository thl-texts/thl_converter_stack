#!env/bin/python
"""
The main Word to Text converter for THL-specific TEI texts. Could be inherited by other types of converters
"""
import os
import logging
import re
import unicodedata
import zipfile
import html
import docx
import json
from lxml import etree

from datetime import date
from styleelements import getStyleElement, fontSame, getFontElement
from w3lib.html import replace_entities

TEMPLATE_FOLDER = 'templates'
IGNORABLE_STYLES = ['Paragraph Char', 'List Bullet Char']
ANNOTATION_PATTERN = r"((?:\*?\s*[A-Z][a-z0-9]+(?:\s+\([0-9\.ab]+\))?,?\s*)+):?\s+" \
                  r"([\u0F00-\u0FFF]+|[oO]mits?|[iI]llegible|[aA]dds|[uU]nclear)"
'''
Explanation of annotation regex above: 
    first parentheses ( = matching group of all sigla with optional pages (most of the first line)
    second parentheses (?: = non-matching group of sigla: [A-Z][a-z0-9] with the above
    third parentheses (?: = non-matching group of optional colon, space, and pagination
                            within parentheses: (?:\s+\([0-9\.ab]+\))? within second parentheses
    fourth parentheses \(...\) = escaped parentheses in third parenthese above part of regex
    fifth parentheses ([\u0F00-\u0FFF] ... ) = matching group of Tibetan reading or keywords (the whole second line)
'''


def get_lang_by_char(chr):
    unm = unicodedata.name(chr)
    if "TIBETAN" in unm:
        return "tib"
    if "CHINESE" in unm:
        return "chi"
    if "DEVANAGARI" in unm:
        return "san"
    return ''


class TextConverter:

    def __init__(self, args):
        self.args = args
        self.files = []
        self.current_file = ''
        self.current_file_path = ''
        self.indir = args.indir
        if not os.path.isdir(self.indir):
            raise NotADirectoryError("The in path, {}, is not a directory".format(self.indir))
        self.getfiles()
        self.outdir = args.out
        self.overwrite = args.overwrite
        if not os.path.isdir(self.outdir):
            raise NotADirectoryError("The out path, {}, is not a directory".format(self.outdir))
        self.metafields = args.metafields if args.metafields else False
        self.template = args.template
        self.xmltemplate = ''
        self.dtdpath = args.dtdpath
        self.debug = args.debug if args.debug else False
        self.worddoc = None
        self.nsmap = None
        self.metatable = None
        self.footnotes = {}
        self.fncount = 0
        self.endnotes = {}
        self.endntcount = 0
        self.xmlroot = None
        self.headstack = []
        self.current_el = None
        self.pindex = -1
        self.edsig = args.edition_sigla
        if self.edsig:
            print(f"Using edition sigla from command argurments: {self.edsig}")
        self.chapnum = None
        self.textid = ''
        self.in_multiline_apparatus = False
        self.multiline_apparatus_num = 0
        self.multiline_apparatus_el = None

        self.log = args.log
        self.loglevel = logging.DEBUG if self.debug else logging.WARN
        logging.basicConfig(level=self.loglevel)

    def getfiles(self):
        files_in_dir = os.listdir(self.indir)
        files_in_dir.sort()
        for sfile in files_in_dir:
            if sfile.endswith(".docx") and not sfile.startswith('~'):
                self.files.append(sfile)

    def setlog(self):
        # fname = os.path.split(fname)[1].replace('.docx', '') + '.log'
        logpath = os.path.join(self.log, self.current_file.replace('docx', 'log'))
        if self.debug:
            print("Log file for {} is: {}".format(self.current_file, logpath))
        loghandler = logging.FileHandler(logpath, 'w')
        log = logging.getLogger()
        for hdlr in log.handlers[:]:
            log.removeHandler(hdlr)
        log.addHandler(loghandler)

    def convert(self):
        for fl in self.files:
            print("\n======================================\nConverting file: {}".format(fl))
            self.current_file = fl
            self.setlog()
            self.convertdoc()
            self.bodydivcheck()
            self.assignids()
            self.tidyxml()
            self.writexml()

    def convertdoc(self):
        self.current_file_path = os.path.join(self.indir, self.current_file)
        self.worddoc = docx.Document(self.current_file_path)
        if self.textid == '':
            mtch = re.search(r"^\S+-\d+-text", self.current_file)
            if mtch:
                self.textid = mtch.group(0)
        self.nsmap = self.worddoc.element.nsmap
        self.merge_runs()
        self.pre_process_notes()
        self.createxml()

        # self.mylog("In self my warning")

        # Iterate through paragraphs
        totalp = len(self.worddoc.paragraphs)
        ct = 0
        in_app = False
        app_ps = []
        for index, p in enumerate(self.worddoc.paragraphs):
            ct += 1
            print("\rDoing paragraph {} of {}  ".format(ct, totalp), end="")
            self.pindex = index
            if isinstance(p, docx.text.paragraph.Paragraph):
                # Checks for and processes multiline apparatus returns true if paragraph is processed
                paragraph_processed = self.process_multiline_app(p)
                # If not in a multiline apparatus, process paragraph normally
                if not paragraph_processed:
                    self.convertpara(p)
            else:
                self.mylog("Warning: paragraph ({}) is not a docx paragraph cannot convert".format(p))
        print("")

    def pre_process_notes(self):
        """
        Preprocess footnotes and endnotes

        TODO: Check how footnote XML deals with italics, bold, etc and convert to TEI compliant

        :return:
        """
        zipdoc = zipfile.ZipFile(self.current_file_path)

        # write content of endnotes.xml into self.footnotes[]
        fntfile = 'word/footnotes.xml'

        if fntfile in zipdoc.namelist():
            fnotestxt = zipdoc.read('word/footnotes.xml')
            xml_fn_root = etree.fromstring(fnotestxt)
            nsmap = xml_fn_root.nsmap  # The MS Word namesapce map for the footnote document
            wns = '{' + nsmap["w"] + '}'  # The string of the particular namespace "w:" used for getting attributes

            # To output the footnote XML file from Word uncomment the lines below:
            # with open('./workspace/logs/footnotes-test.xml', 'wb') as xfout:
            #     xfout.write(etree.tostring(xml_fn_root))

            fnotes = xml_fn_root.findall('w:footnote', nsmap)
            for fnindex, f in enumerate(fnotes):
                if fnindex > 1:  # The first two "footnotes" are the separation and continuation lines
                    # Footnote object saved in footnote dictionary of class
                    fno = {
                        'num': '',
                        'is_annotation': False,
                        'ref': None,
                        'prev_el': None,
                        'runs': None,
                        'text': '',
                        'markup': ''
                    }
                    if len(f.keys()) > 0:
                        fnum = f.get(f'{wns}id')
                        fno['ref'] = self.worddoc.element.xpath(f'//w:footnoteReference[@w:id="{fnum}"]')
                        fno['num'] = fnum
                        if len(fno['ref']) > 0:
                            fno['ref'] = fno['ref'][0]
                            fno['prev_el'] = fno['ref'].getparent().getprevious()
                            loopct = 0
                            while fno['prev_el'] is not None and fno['prev_el'].tag != f'{wns}r' and loopct < 20:
                                loopct += 1
                                fno['prev_el'] = fno['prev_el'].getprevious()
                            fno['prev_run'] = self.get_run_before_note(fnum)

                    # All runs in footnote. Footnote is a single wrapper element the "r" elements are runs
                    fno['runs'] = f[0].findall("w:r", nsmap)
                    # Iterate through runs in the footnote
                    for fnr in fno['runs']:
                        # Find all the text elements in the run (usually only 1, but this is just in case)
                        for fntxt in fnr.findall("w:t", nsmap):
                            fno['text'] += fntxt.text  # add to the fnot plain text property
                            # See if there is a previous Word Style element with a style name for this run
                            prevel = fntxt.getprevious()
                            if prevel is not None:
                                stylel = prevel.findall('w:rStyle', nsmap)
                                # If there is get the style name from it val attribute
                                if len(stylel) > 0:
                                    stylel = stylel[0]
                                    stylename = stylel.get(f'{wns}val')
                                    stylel = getStyleElement(stylename)
                                    # If it's a legit style, create a element with this text
                                    if stylel is not None:
                                        stylel.text = fntxt.text
                                        # Convert element to string and add to markup property
                                        fno['markup'] += etree.tostring(stylel).decode('utf-8')
                                    else:
                                        # If not a legit style, just add the text of the run, diretctly to the markup;
                                        fno['markup'] += fntxt.text
                                else:  # This led to just a dash in the markup whereas text was more robust
                                    fno['markup'] += fntxt.text
                    # Warn if we can't find the footnote reference number
                    if not fno['num']:
                        pretext = fno['text'][:25]
                        print(f"\n\tNo footnote number found for note index {fnindex}, beginning with “{pretext}”")
                    else:
                        fno['is_annotation'] = self.fn_is_annotation(fno)
                        fnkey = fno['num']
                        if type(fnkey) == str and type(self.footnotes) == dict:
                            self.footnotes[fnkey] = fno

                    # Old Footnote code for reference (july 29, 2022)
                    # s = ""
                    # plains = ""
                    # wdschema = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                    # for t in text:
                    #     ttxt = t.text
                    #     plains += ttxt
                    #     prev = t.getprevious()  # This returns <w:rPr> or None
                    #     if prev is not None:
                    #         tsty = []
                    #         lang = ""
                    #         for pc in prev.getchildren():
                    #             pcstyle = re.sub(r'\{[^\}]+\}', '', pc.tag)
                    #             if pcstyle == 'rStyle':
                    #                 if pc.get(wdschema + 'val') == 'X-EmphasisStrong':
                    #                     tsty.append('strong')
                    #             elif pcstyle == 'i':
                    #                 tsty.append('weak')
                    #             elif pcstyle == 'u':
                    #                 tsty.append('underline')
                    #             elif pcstyle == 'lang':
                    #                 if pc.get(wdschema + 'bidi') == 'bo-CN':
                    #                     lang = ' lang="tib"'
                    #         attr = "" if len(tsty) == 0 else ' rend="{}"'.format(' '.join(tsty))
                    #         if lang == "":
                    #             langcode = get_lang_by_char(ttxt[0])
                    #             if len(langcode) > 0:
                    #                 lang = ' lang="{}"'.format(langcode)
                    #         attr += lang
                    #         ttxt = '<hi{}>{}</hi>'.format(attr, ttxt) if len(attr) > 0 else ttxt
                    #     s += ttxt
                    # TODO: add regex here to find <hi> with the same rend next to each other and merge them,
                    #  e.g. <hi lang="tib">ཡོད</hi><hi lang="tib">།</hi>
                    # Old footnote markup for refence
                    # nteltxt = f'<note type="footnote" n="{fn_num}">{s}<rs>{plains}</rs></note>'
                    # try:
                    #     note_el = etree.XML(nteltxt.replace('&', '&amp;'))  # convert & to its xml entity
                    #     self.footnotes[fn_num] = note_el
                    # except etree.XMLSyntaxError as xe:
                    #     # Information if there is a problem creating XML
                    #     print("Xml syntax error in creating note: ")
                    #     print(xe)
                    #     print(f"nt el: {nteltxt}")
                    #     print(f"s: {s}")
                    #     print(f"plains: {plains}")
                # fnindex += 1

        endntfile = 'word/endnotes.xml'
        if endntfile in zipdoc.namelist():
            # write content of endnotes.xml into self.endnotes[]
            xml_content = zipdoc.read(endntfile)
            xml_en_root = etree.fromstring(xml_content)

            # To output the endnote XML file from Word uncomment the lines below:
            with open('./workspace/logs/endnotes-test.xml', 'wb') as xfout:
                xfout.write(etree.tostring(xml_en_root))

            enindex = 0
            enotes = xml_en_root.findall('w:endnote', xml_en_root.nsmap)
            if len(enotes) > 2:
                self.mylog("There are endnotes! Need to update endnote processing code!")

            for f in enotes:
                if enindex > 1:
                    text = f.findall('.//w:t', xml_en_root.nsmap)
                    s = ""
                    wdschema = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                    for t in text:
                        ttxt = t.text
                        prev = t.getprevious()  # This returns <w:rPr> or None
                        if prev is not None:
                            tsty = []
                            for pc in prev.getchildren():
                                pcstyle = re.sub(r'\{[^\}]+\}', '', pc.tag)
                                if pcstyle == 'rStyle':
                                    if pc.get(wdschema + 'val') == 'X-EmphasisStrong':
                                        tsty.append('strong')
                                if pcstyle == 'i':
                                    tsty.append('weak')
                                if pcstyle == 'u':
                                    tsty.append('underline')
                            ttxt = '<hi rend="{}">{}</hi>'.format(' '.join(tsty), ttxt)
                        s += ttxt
                    note_el = etree.XML('<note type="endnote">{}</note>'.format(s))
                    self.endnotes.append(note_el)
                enindex += 1
        zipdoc.close()
        self.fncount = len(self.footnotes)
        self.endntcount = len(self.endnotes)

    def createxml(self):
        template_path = os.path.join(TEMPLATE_FOLDER, self.template)
        with open(template_path, 'r') as tempstream:
            if self.debug:
                self.mylog(f"Template file: {template_path}")
            self.xmltemplate = tempstream.read()
            self.metatable = self.worddoc.tables[0] if len(self.worddoc.tables) else False
            if self.metatable:
                self.createmeta()  # Separated out to be overriden. Creates the XML string with metadata inserted
            # Take self.xmltemplate the xml string and convert to an XML object
            xmldoc = bytes(bytearray(self.xmltemplate, encoding='utf-8'))
            # create lxml element tree from metadata info
            parser = etree.XMLParser(ns_clean=True, recover=True, encoding='utf-8')
            self.xmlroot = etree.fromstring(xmldoc, parser)

    def createmeta(self):
        """
        This table creates the xmlstring by filling in the template string with the table values based on the
        table labels. It uses the classes self.metatable property which is the Word docs metatable and
        self.template which is the string read in from the XML template file indicated in the settings

        NOTE: This function can be overridden by particular template converters to customize

        :return:
        """
        wordtable = self.metatable
        # Fill out metadata matching on string in wordtable with {strings} in template (teiHeader.dat)
        xmltext = self.xmltemplate.replace("{Digital Creation Date}", str(date.today()))
        problems_on = False
        tablerows = len(wordtable.rows)
        problems = []
        print("Process Metadata Table ...")
        for rwnum in range(0, tablerows):
            try:
                if wordtable._column_count < 2:
                    label = wordtable.cell(rwnum, 0).text.strip()
                    logging.warning("Row {} of metadata table has too few cells".format(rwnum))
                    continue

                elif wordtable._column_count == 2:
                    label = wordtable.cell(rwnum, 0).text.strip()
                    rowval = wordtable.cell(rwnum, 1).text.strip()

                elif wordtable._column_count == 3:
                    label = wordtable.cell(rwnum, 0).text.strip()
                    rowval = wordtable.cell(rwnum, 1).text.strip()
                    self.mylog("Notice: Not handling third column of row {}!".format(rwnum))

                elif wordtable._column_count == 4 and wordtable.cell(rwnum, 0).text == wordtable.cell(rwnum, 2).text:
                    # Not giving warning because this seems to be a Word or pydoc weirdness
                    # In the stylized metadata table, it detects four columns per row
                    # with the first 3 all containing the same label. So just using last two.
                    label = wordtable.cell(rwnum, 2).text.strip()
                    rowval = wordtable.cell(rwnum, 3).text.strip()

                else:
                    rowcells = wordtable.rows[rwnum].cells
                    label = wordtable.cell(rwnum, 0).text.strip() if len(rowcells) > 0 else ""
                    rowval = wordtable.cell(rwnum, 1).text.strip() if len(rowcells) > 1 else label
                    # self.mylog("Row {} of metadata table has too many ({}) cells. Using first two".format(rwnum))

                if label == "Text ID":
                    self.textid = rowval
                # print(label)
                # All Uppercase are Headers in the table skip
                if label.isupper():
                    continue  # Skip labels
                if 'problems' in label.lower():
                    label = "Problems"  # Must be "{Problems}" in template
                    rowval = self.getTextFromCell(wordtable.cell(rwnum, 1))
                    if rowval is None or len(rowval) == 0:
                        rowval = "No problems"

                temppt = label.split(' (')  # some rows have " (if applicable)" or possible some other instruction

                # Normalizing Label
                label = temppt[0].replace('\xa0', ' ').strip()
                label = label.replace(' (if applicable)', '')
                label = label.replace('Call་number', 'Call-number')
                # 'Cover Title Tib' is in XML comment while 'Title on Cover' is in form (maybe a later correction)
                if label in ('Title on Cover', 'Title on Cover Page', 'Title on Title Page'):
                    label = 'Cover Title Tib'
                # Normalize 'Cover Page Title' and 'Title Page Title' to 'Cover Title' in all occurrences
                if label != 'Cover Page':
                    label = label.replace('Cover Page', 'Cover').replace('Title Page', 'Cover')
                srclbl = "{" + label + "}"

                # Replace occurence of label in XML header
                xmltext = xmltext.replace(srclbl, rowval)

                # Deal with edition sigla
                if label.lower() == 'edition sigla':
                    self.edsig = rowval

                # Deal with Chapter Number
                if label.lower() == 'chapter number':
                    self.chapnum = rowval

            except IndexError as e:
                logging.error("Index error in iterating wordtable: {}".format(e))
            except TypeError as e:
                logging.error("Type error in iterating wordtable: {}".format(e))

        # Add text ID if necessary and in current file name
        if self.current_file:
            res = re.search(r'^(\w+-\d+)-text', self.current_file)
            if res:
                xmltext = xmltext.replace('{Text ID}', res.group(1))

        self.xmltemplate = re.sub(r'{([^}]+)}', r'<!--\1-->', xmltext)  # comment out any unreplaced labels

    def convertpara(self, p):
        style_name = p.style.name
        headmtch = re.match(r'^Heading (?:Tibetan\s*)?(\d+)[\,\s]*(Front|Body|Back)?', style_name)
        if headmtch:
            self.do_header(p, headmtch)

        elif len(self.headstack) == 0:
            # if there is not yet a headstack then it's notes at beginning of document that should be ignored
            ptxt = p.text[0:150] if len(p.text) > 150 else p.text
            self.mylog(f"Skipping Paragraph at beginning: {ptxt}")
            return

        elif "List" in style_name:
            self.do_list(p)

        elif "Verse" in style_name:
            self.do_verse(p)  # Note this does verse citation and verse speech as well

        elif "Citation" in style_name:
            self.do_citation(p)  # Note verse citation is done above in do_verse()

        elif "Section" in style_name:
            doruns = self.do_section(p)
            if not doruns:
                return

        elif "Speech" in style_name:
            self.do_speech(p)  # verse speech is done in do_verse()

        else:
            if not self.is_reg_p(style_name):
                msg = "\n\tStyle {} defaulting to paragraph".format(style_name)
                self.mylog(msg)
            self.reset_current_el()
            self.do_paragraph(p)

        # Once paragraphs have been processed. self.current_el is the element where the runs of the paragraph go
        self.iterate_runs(p)  # so all we need to do is send the word paragraph object

    def do_header(self, p, headmtch):
        """
        Method to convert heading styles into structural elements either front, body, bock or divs
        Called from convertpara() method above

        :param p: the word docx paragraph object
        :param headmtch: the re.match object group(1) is heading level, group(2) is front, body, back if it is one of those
        :return: none
        """
        hlevel = int(headmtch.group(1))
        style_name = p.style.name
        if hlevel == 0:
            # If level is 0, its front body or back, create element and clear head stack
            fbbel = etree.XML('<{0}><head></head></{0}>'.format(headmtch.group(2).lower())) # the match is e.g. "front"
            self.xmlroot.find('text').append(fbbel)
            self.current_el = fbbel.find('head')
            self.headstack = [fbbel]
        else:
            # Otherwise we are already in front, body, or back, so create div
            currlevel = len(self.headstack) - 1  # subtract one bec. div 0 is at top of stack
            hdiv = etree.XML(f'<div n="{hlevel}"><head></head></div>')
            # if it's the next level deeper
            if hlevel > currlevel:
                # if new level is higher than the current level just add it to current
                if hlevel - currlevel > 1:
                    self.mylog("Warning: Heading level skipped for {}".format(style_name, p.text))
                self.headstack[-1].append(hdiv)  # append the hdiv to the previous one
                self.headstack.append(hdiv)      # add the hdiv to the stack
            # if it's the same level as current
            elif hlevel == currlevel:
                if len(self.headstack) > 0:
                    self.headstack[-1].addnext(hdiv)
                    self.headstack[-1] = hdiv
                else:
                    errmsg = f"Headstack is empty when adding div ({hdiv.text})\n"
                    errmsg += "Make sure Body Heading0 is present."
                    raise ConversionException(errmsg)

            # Otherwise it's a higher level
            else:
                # because front, body, back is 0th element use hlevel to splice array
                self.headstack = self.headstack[0:hlevel]  # splice arra to parent elements
                self.headstack[-1].append(hdiv)
                self.headstack.append(hdiv)
            self.current_el = hdiv.find('head')

    def do_list(self, p):
        # Get current and previous list styles and numbers
        my_style = p.style.name
        my_num = self.getnumber(my_style) or 1
        prev_style = self.get_previous_p(True)
        prev_num = self.getnumber(prev_style) if "List" in prev_style else False
        if not prev_num:
            if my_style == prev_style:
                prev_num = my_num
            elif "List" in prev_style:
                prev_num = 1

        # Determine the type of list and set attribute values
        rend = "bullet" if "Bullet" in my_style else "1"
        nval = ' n="1"' if rend == "1" else ""
        # ptxt = p.text

        # Create element object templates (used below in different contexts). Not all are used in all cases
        itemlistel = etree.XML('<item><list rend="{}"{}><item></item></list></item>'.format(rend, nval))
        listel = etree.XML('<list rend="{}"{}><item></item></list>'.format(rend, nval))
        itemel = etree.XML('<item></item>')

        # Different Alternatives
        if not prev_num:  # new non-embeded list
            if my_num and int(my_num) > 1:
                self.mylog("List level {} added when not in list".format(my_num))
            # Add list after current element and make it current
            self.current_el.addnext(listel)
            self.current_el = listel.find('item')

        else:   # For Lists embeded in Lists
            my_num = int(my_num)
            prev_num = int(prev_num)
            if my_num > prev_num:  # Adding/Embedding a new level of list
                if my_num > prev_num + 1:
                    self.mylog("Skipping list Level: inserting {} at level {}".format(my_num, prev_num))
                # Append the itemlist el to the current list
                self.current_el.addnext(itemlistel)
                self.current_el = itemlistel.find('list/item')

            elif my_num == prev_num: # On same level lists
                self.current_el.addnext(itemel)
                self.current_el = itemel

            elif self.current_el is not None:  # my_num < prev_num  Returning to an ancestor list with lower list number
                lvl = prev_num
                listel = self.current_el.getparent()
                try:
                    while lvl != my_num:
                        listel = listel.getparent()
                        if listel.tag == 'list':
                            lvl -= 1
                except AttributeError:
                    self.mylog("No iterancestors() method for current list element (text: {}). " +
                                   "Current Element Class: {}".format(p.text, self.current_el.__class__.__name__))
                listel.append(itemel)
                self.current_el = itemel
            else:  # No current el so use the last div
                self.headstack[-1].append(listel)
                self.current_el = listel.find('item')

    def do_verse(self, p):
        # TODO: Throw warning when a Verse2 is found without a preceding verse1 (maybe interpret it as verse1)
        my_style = p.style.name
        prev_style = self.get_previous_p(True)  # TODO: Check if current style is same (except number) with previous
        is_cite = True if 'citation' in my_style.lower() else False
        is_speech = True if 'speech' in my_style.lower() else False
        is_nested = True if 'nested' in my_style.lower() else False
        is_same = True if my_style.lower().replace('2', '1') == prev_style.lower().replace('2', '1') else False
        level = 2 if '2' in my_style else 1
        if level == 2:
            myel = etree.Element('l')
            if not is_nested and "nested" in prev_style.lower():
                self.current_el.getparent().addnext(myel)
            else:
                self.current_el.addnext(myel)
            self.current_el = myel

        elif is_cite:
            if is_same:
                markup = etree.XML('<lg><l></l></lg>')
                self.current_el.getparent().addnext(markup)
                self.current_el = markup.find('l')

            elif is_nested and "citation" in prev_style.lower() and "paragraph" in prev_style.lower():
                # when nested in paragraph citation
                markup = etree.XML('<lg><l></l></lg>')
                self.current_el.addnext(markup)
                self.current_el = markup.find('l')

            else:
                markup = etree.XML('<quote><lg><l></l></lg></quote>')
                if is_nested:
                    self.current_el.append(markup)
                else:
                    self.headstack[-1].append(markup)
                self.current_el = markup.find('lg').find('l')

        elif is_nested:
            markup = etree.XML('<lg><l></l></lg>')
            self.current_el.addnext(markup)  # if nested, current el is "l", add "lg" next to this.
            self.current_el = markup.find('l')

        else:
            if is_speech:
                markup = etree.XML('<q><lg><l></l></lg></q>')
                self.current_el = markup.find('lg').find('l')
            else:
                markup = etree.XML('<lg><l></l></lg>')
                self.current_el = markup.find('l')
            self.headstack[-1].append(markup)

    def do_citation(self, p):
        my_style = p.style.name
        prev_style = self.get_previous_p(True)  # TODO: Check if current style is same (except number) with previous
        nested = True if "nested" in my_style.lower() else False
        continued = True if "continued" in my_style.lower() else False

        if continued or (nested and my_style == prev_style):
            cite_el = etree.XML('<p rend="cont"></p>')
            if "verse" in prev_style.lower():
                self.current_el.getparent().addnext(cite_el)
            else:
                self.current_el.addnext(cite_el)
            self.current_el = cite_el
        else:
            cite_el = etree.XML('<quote><p></p></quote>')
            if nested:
                self.current_el.append(cite_el)
            elif "nested" in prev_style:
                self.current_el.getparent().addnext(cite_el)
            else:
                self.current_el.addnext(cite_el)
            self.current_el = cite_el.find('p')

    def do_section(self, p):
        my_style = p.style.name
        ptext = p.text
        nmtch = re.match(r'section\s+(\d)', my_style, re.IGNORECASE)
        if nmtch:
            n = nmtch.group(1)
            sect_el = etree.XML('<milestone unit="section" n="{}" rend="{}" />'.format(n, ptext))
            self.headstack[-1].append(sect_el)
            self.current_el = sect_el
            return False

        elif 'chapter element' in my_style.lower():
            sect_el = etree.XML('<milestone unit="section" n="cle" rend="{}" />'.format(ptext))
            self.headstack[-1].append(sect_el)
            self.current_el = sect_el
            return False

        elif 'interstitial' in my_style.lower():
            sect_el = etree.XML('<div type="interstitial"><head></head></div>')
            self.headstack[-1].append(sect_el)
            self.current_el = sect_el.find('head')
            return True

        else:
            logging.warning("Unknown section type with header: {}".format(ptext))
            # TODO: Should this be a milestone instead of a div???
            sect_el = etree.XML('<div type="section"><head></head></div>')
            self.headstack[-1].append(sect_el)
            self.current_el = sect_el.find('head')
            return True

    def do_speech(self, p):
        #  Note this is speech not already covered in verse or citation. See convertpara() method above
        my_style = p.style.name.lower()
        prev_style = self.get_previous_p(True)
        iscont = True if "continued" in my_style else False
        isnested = True if "nested" in my_style else False
        speech_el = etree.XML('<q><p></p></q>')  # use for new nested or new not nested
        if iscont:
            # iscontinued whether nested or not
            speech_el = etree.XML('<p rend="cont"></p>')
            if "nested" in prev_style and "nested" not in my_style:
                self.current_el.getparent().addnext(speech_el)
            else:
                self.current_el.addnext(speech_el)
            self.current_el = speech_el

        elif isnested:
            # initial nested speech, then the current el is a <p> within a <q> and another <q> after it
            self.current_el.addnext(speech_el)
            self.current_el = speech_el.find('p')

        else:
            # Regular "Speech Paragraph" style
            # Get out from within verse
            if self.current_el.tag == 'l':
                lgs = [lg for lg in self.current_el.iterancestors('lg')]
                self.current_el = lgs[-1] if len(lgs) > 0 else self.current_el.getparent()
            # Get out from within list
            if self.current_el.tag == 'item':
                lists = [lst for lst in self.current_el.iterancestors('list')]
                self.current_el = lists[-1] if len(lists) > 0 else self.current_el.getparent()
            # Add <q><p>... as sibling of current element.
            self.current_el.addnext(speech_el)
            self.current_el = speech_el.find('p')

    def do_paragraph(self, p):
        if p.text.strip() == '':
            return
        p_el = etree.Element('p')
        # TODO: Shouldn't we just append to last element in headstack? What happens after embedded lists?
        if self.current_el is not None:
            if self.current_el.tag == 'div':
                self.current_el.append(p_el)
            else:
                self.current_el.addnext(p_el)
        self.current_el = p_el

    def iterate_runs(self, p, skip=0):
        '''
        Populates a paragraph level element with its runs properly marked up (these are character level styles)
        Creates a <temp> element to contain the inner XML structure of the paragraph level element
        Iterates over the runs in the paragraph creating "elem" elements sometimes with inner structure.
        These get appended to <temp> which in the end is added to the doc and becomes self.current_el

        [In old converter this was iterateRange (the interateRuns function was not called)]

        :param p:
        :param skip: (int) number of runs to skip before beginning processing (used for multiline apparatus)
        :return:
        '''
        last_run_style = ''
        style_before_footnote = ''
        temp_el = etree.Element('temp')  # temp element to put xml element objects in
        elem = None
        if len(p.runs) == 0:
            temp_el.text = " "
            return

        if temp_el.text is None:
            temp_el.text = ""

        lemma_contents = []
        for rct, run in enumerate(p.runs):
            if run is None or rct < skip:
                continue
            rtxt = run.text
            if elem is not None and elem.text is None:
                elem.text = ""
            if elem is not None and elem.tail is None:
                elem.tail = ""

            if "{" in rtxt and "}" not in rtxt:
                # self.mylog(f"Open but no close brace: {rtxt}")
                lemma_contents.append(run)
                nextct = rct
                while "}" not in p.runs[nextct].text and nextct - rct < 100:
                    nextct += 1
                    lemma_contents.append(p.runs[nextct])
                if "}" in p.runs[nextct].text and p.runs[nextct + 1].style.name == 'footnote reference':
                    skip = nextct + 1
                    if rtxt == '{' or rtxt[0] == '{':
                        continue
                    elif rtxt[-1] == '{':
                        lemma_contents.pop(0)
                        rtxt = rtxt[:-1]
                        run.text = rtxt
                    else:
                        pts = rtxt.split('{')
                        rtxt = pts[0]
                        run.text = rtxt
                        lemma_contents[0] = pts[1]
                else:
                    self.mylog(f"Did not find closing brace for initial opening brace. Ignoring: {rtxt}")
                    lemma_contents = []


            # if "Heading" in p.style.name:
            #    rtxt = re.sub(r'^[\d\s\.]+', '', rtxt)

            char_style = run.style.name
            is_new_style = True if char_style != last_run_style else False
            # if "root" in char_style.lower():
            #    print(f"root style: {char_style}")

            # Default Paragraph Font
            if not char_style or char_style == "" or "Default Paragraph Font" in char_style:
                # May be in Default Font but have bold or italic ste
                new_el = getFontElement(run)  # Check for those font characteristics
                if new_el is not None:
                    new_el.text = rtxt
                    temp_el.append(new_el)
                    elem = temp_el.getchildren()[-1]
                elif elem is None:
                    temp_el.text += rtxt
                else:
                    elem.tail += rtxt
            # Endnotes
            elif "endnote" in char_style.lower():
                # TODO: Deal with Endnotes
                self.mylog("\n\tNEED TO DEAL WITH ENDNOTES!!!")

            # Footnotes
            elif "footnote" in char_style.lower():
                style_before_footnote = last_run_style
                fnnum, note = self.getFootnoteFromRefRun(run)
                if not note:
                    self.mylog(f"\n\tCould not find footnote object for {fnnum}")
                    return

                if note['is_annotation']:
                    if elem is None and len(temp_el.getchildren()) > 0:
                        elem = temp_el.getchildren()[-1]
                    if elem is not None:
                        back_text = elem.tail if elem.tail and len(elem.tail) > 0 else elem.text
                    else:
                        back_text = temp_el.text
                    # print(f"back text: {back_text}")
                    reading = self.process_critical(note, back_text, lemma_contents)
                    lemma_contents = []  # reset lemma contents list after being processed
                    if not reading:
                        self.mylog("no reading discovered!")
                        self.mylog(f"{note['num']}, {note['text']}, {note['markup']}")
                    elif isinstance(elem, etree._Element):
                        if elem.tag != 'milestone' and (elem.tail is None or len(elem.tail) == 0):
                            elem.text = reading['backtext']
                            elem.append(reading['app'])
                        else:
                            elem.tail = reading['backtext']
                            temp_el.append(reading['app'])
                            elem = reading['app']
                    else:
                        temp_el.text = reading['backtext']
                        temp_el.append(reading['app'])
                        elem = reading['app']
                else:   # Other note is not annotation but regular text
                    note_mu = note['markup'] if note['markup'] else note['text']
                    # if len(note['markup']) < len(note['text']):
                    #    note_mu = note['text']
                    try:
                        note_mu = html.escape(note_mu, True)
                        note_mu = etree.XML(f'<note type="footnote" n="{fnnum}">{note_mu}</note>')
                        temp_el.append(note_mu)
                        elem = note_mu
                    except etree.XMLSyntaxError as XSE:
                        self.mylog("XML Syntax Error: {}".format(XSE))
                        self.mylog("On note {}, with markup: {}".format(fnnum, note_mu))

            # Milestones
            elif "Page Number" in char_style or "Line Number" in char_style:
                # If there are multiple ms of the same style, they get merged in merge_runs. So split them up
                msitems = rtxt.split('][')
                for mstxt in msitems:
                    elem = self.createmilestone(char_style, mstxt)
                    temp_el.append(elem)

            elif is_new_style or elem is None or elem is False:
                new_el = getStyleElement(char_style)
                if new_el is None:
                    temp_el.text += rtxt
                    if self.debug and char_style not in IGNORABLE_STYLES:
                        outrtxt = rtxt.strip()
                        pstart = p.runs[0].text if len(p.runs) > 0 else p.text
                        if len(pstart) > 25:
                            pstart = pstart[0:25]
                        msg = f"\n\tNo style definition found for style name, {char_style}: {outrtxt}\n" \
                              f"\tParagraph beginning with: “{pstart}”..."
                        self.mylog(msg)

                else:
                    if self.debug:
                        logging.debug('Style element {} => {}'.format(char_style, new_el.tag))
                    new_el.text = rtxt
                    temp_el.append(new_el)
                    elem = temp_el.getchildren()[-1]
            else:
                elem.text += rtxt

            last_run_style = char_style

        ### End of iterating runs ###

        if self.current_el is None:
            self.mylog('current el is none') # Should never get here

        # Copy temp_el contents to current_el depending on whether it has children or not
        if self.current_el.tag in ['anchor', 'lb', 'pb', 'milestone']:
            empty_el = self.current_el
            for tmpchld in temp_el.getchildren():
                self.current_el.addnext(tmpchld)
                self.current_el = tmpchld
            empty_el.tail = temp_el.text
            self.current_el = self.current_el.getparent()
            return  # No need to do the following

        curr_child = list(self.current_el) or []
        if len(curr_child) > 0:
            curr_child[-1].tail = temp_el.text
        else:
            self.current_el.text = temp_el.text
        for tempchild in list(temp_el):
            self.current_el.append(tempchild)

        # Deal with numbers at the beginning of headers
        if "heading" in p.style.name.lower():
            headtxt = self.current_el.text
            mtch = re.match(r'^((\d+\.?)+)', headtxt)
            if mtch:
                hnum = mtch.group(1)
                headtxt = headtxt.replace(hnum, '')
                numel = etree.XML('<num>{}</num>'.format(hnum))
                numel.tail = headtxt
                self.current_el.text = ""
                self.current_el.insert(0, numel)

    def process_critical(self, note, bcktxt, lemma_contents):
        '''
        Processes footnotes for critical edition with alternate readings from other editions

        :param note: the note object created for this not during pre_process_notes
        :param bcktxt: the text preceding the note reference containing the lemma in braces
        :return:
        '''
        reading = {}

        if bcktxt is None or len(bcktxt) == 0:
            reading['backtext'] = ''
            lem = ''
        elif bcktxt[-1] == '}':
            cestind = bcktxt.rfind('{')
            lem = bcktxt[cestind + 1:len(bcktxt) - 1].strip() if cestind > -1 else "FIX"
            reading['backtext'] = bcktxt[0:cestind]
        elif len(lemma_contents) > 0:
            reading['backtext'] = ""
        elif len(bcktxt) == 1:
            lem = bcktxt
            reading['backtext'] = ""
        else:
            syls = re.split(r"\u0F0B", bcktxt)
            lem = syls.pop()  # when preceding text ends in tsek this will be empty
            if lem == '':
                lem = syls.pop() + "\u0F0B"
            reading['backtext'] = "\u0F0B".join(syls) + "\u0F0B"

        lemed = self.edsig if self.edsig and self.edsig != '' else 'base'
        lempg = ''
        notedata = TextConverter.process_critical_note(note)
        if notedata['lem']:
            lemed = notedata['lem']['edsigs']
            lempg = ' n=""'.format(notedata['lem']['edpgs'])
        ntnum = note['num']

        # If there's marked up content in the lemma braces, get the XML string for the lem value
        if len(lemma_contents) > 0:
            lem = self.process_lemma_contents(lemma_contents)

        # Check for readings that are just a single punctuation mark, if one exists
        # Then trim lemma to last character (punctuation) and put extra in the backtext
        if len(lem) > 1:
            # RegEx for a single punctuation mark within the range between tsek and gter shad
            tibpunct ='[\u0F0B-\u0F14]'
            for vrnt in notedata['variants']:
                if re.match(tibpunct, vrnt['txt']):
                    reading['backtext'] += lem[0:-1]
                    lem = lem[-1]
                    break

        app = f'<app n="nt{ntnum}" id="app{ntnum}"><lem wit="{lemed}"{lempg}>{lem}</lem>'
        for vrnt in notedata['variants']:
            natt = ''
            edpgs = vrnt['edpgs']
            edsigs = vrnt['edsigs']
            if len(vrnt['edpgs'].strip(' ')) > 0:
                natt = f' n="{edpgs}"'
            if vrnt['pref']:
                natt += ' rend="pref"'
            vartxt = lem if vrnt['txt'] == '' else vrnt['txt']
            # For English descriptive terms in annotation, make into type attribute and use no text
            keywords = ['omit', 'illegible', 'unclear', 'corrupt']
            if any([word in vartxt.lower() for word in keywords]):
                vartxt = vartxt.lower()
                if vartxt == 'omits':
                    vartxt = 'omit'
                natt += f' lang="eng" type="{vartxt}"'
                vartxt = ''

            app += f'<rdg wit="{edsigs}"{natt}>{vartxt}</rdg>'
        if notedata['note']:
            app += f'<wit rend="note">{notedata["note"]}</wit>'
        if len(notedata['interp']) > 0:
            app += f'<interp value="{notedata["interp"]}"/>'
        app += '</app>'
        reading['app'] = etree.XML(app)
        if lem == "FIX":
            ntnumb = note['num']
            nteltxt = note['text']
            self.mylog(f"\n\t“FIX” Footnote {ntnumb} follows close brace as for apparatus, "
                           f"but no preceding open brace detected: "
                           f"\n\tWord before note: “{bcktxt}”"
                           f"\n\tNote: {nteltxt}")
        return reading

    def process_lemma_contents(self, lcnts):
        lemmaout = ""
        for rn in lcnts:
            if isinstance(rn, str):
                lemmaout += rn
            elif isinstance(rn, bytes):
                lemmaout += rn.decode('utf-8')
            elif isinstance(rn, docx.text.run.Run):
                rstyle = rn.style.name
                rtxt = rn.text
                if "Page Number" in rstyle or "Line Number" in rstyle:
                    # If there are multiple ms of the same style, they get merged in merge_runs. So split them up
                    msitems = rtxt.split('][')
                    for mstxt in msitems:
                        elem = self.createmilestone(rstyle, mstxt)
                        lemmaout += etree.tostring(elem).decode('utf-8')
                else:
                    new_el = getStyleElement(rstyle)
                    if new_el is None:
                        lemmaout += rtxt
                    else:
                        new_el.text = rtxt
                        lemmaout += etree.tostring(new_el).decode('utf-8')
        return lemmaout.replace('{', '').replace('}', '')

    @staticmethod
    def process_critical_note(anote):
        notedata = {
            'lem': False,
            'variants': [],
            'note': False,
            'interp': ''
        }

        appnote = re.search(r'\[[^\]]+\]', anote['text'])
        if appnote is not None:
            notedata['note'] = appnote.group(0)
            anote['text'] = anote['text'].split(notedata['note'])[0].strip()
            notedata['note'] = notedata['note'].strip('[] ')


        # notepts = anote['text'].strip(' .').split(';')
        notepts = anote['text'].split('. ', 1)
        if len(notepts) > 1 and notepts[1] != '':
            notedata['interp'] = notepts[1]  # Peel off the descriptive note
        notepts = notepts[0].strip(' .').split(';')

        # Old pattern for separating reading from edition sigla, see ANNOTATION_PATTERN above
        # pattern = r"((?:[A-Z][a-z0-9]+,?\s*)+):?\s+([\u0F00-\u0FFF]+|[oO]mits|[iI]llegible|[aA]dds|[uU]nclear)"

        for rdg in notepts:
            pref = True if '*' in rdg else False
            rdg = rdg.strip(' *')
            mtc = re.search(ANNOTATION_PATTERN, rdg)
            if mtc:
                rpts = [mtc.group(1), mtc.group(2)]
            else:
                rpts = rdg.split(':')
            tempeds = [ed.strip() for ed in rpts[0].split(',')]
            edsigs = []
            edpgs = []
            for ed in tempeds:
                epts = ed.replace(')', '').split('(')
                edsigs.append(epts[0])
                edpg = epts[1] if len(epts) > 1 else ''
                edpgs.append(edpg)

            edsigs = ' '.join(edsigs)
            edpgs = ' '.join(edpgs)

            # if editions listed without a reading text, they are the lemma sources
            rdgtxt = rpts[1] if len(rpts) > 1 else False
            if not rdgtxt:
                notedata['lem'] = {
                    'edsigs': edsigs,
                    'edpgs': edpgs
                }
            else:
                notedata['variants'].append({
                    'pref': pref,
                    'edsigs': edsigs,
                    'edpgs': edpgs,
                    'txt': rdgtxt.strip() if rdgtxt else ''
                })
        return notedata

    def process_multiline_app(self, p):
        """
        Processes multi paragraph apparatus by surrounding them in two empty tags with corresponding IDs:
                <addSpan id="span1-open" to="span1-close" rend="apparatus" n="Dg Ab"/>
                  (More markup here)
                <anchor id="span1-close"  corresp="span1-open" rend="apparatus"/>

        :param p:
        :return:
        """
        paragraph_processed = False
        ptxt = p.text
        # Detect if there's a multiline apparatus and begin the processing
        if len(ptxt) > 0 and ptxt[0] == '{' and '}' not in ptxt:
            self.mylog('Multiline apparatus begins: ' + ptxt)
            p.runs[0].text = p.runs[0].text[1:]
            self.in_multiline_apparatus = True
            self.multiline_apparatus_num += 1
            self.multiline_apparatus_el = etree.XML(f'<addSpan id="span{self.multiline_apparatus_num}-open"'
                                                    f' rend="apparatus" ></addSpan>')
            if self.current_el is not None:
                if self.current_el.tag == 'div':
                    self.current_el.append(self.multiline_apparatus_el)
                else:
                    self.current_el.addnext(self.multiline_apparatus_el)
            self.current_el = self.multiline_apparatus_el

        # Process paragraphs within a multiline appratus
        elif self.in_multiline_apparatus:
            if ptxt[0] == '}':  # closing brace must be first character in line
                self.mylog("Multiline apparatus finished: " + ptxt)
                paragraph_processed = True  # This returns true to prevent further processing of this paragraph
                srcs = []
                if p.runs[1].style.name == 'footnote reference':
                    nnum, note = self.getFootnoteFromRefRun(p.runs[1])
                    srcs = [pt.strip() for pt in note['text'].split(',')]
                else:
                    self.mylog("Closing brace for multi-line apparatus is not followed by footnote")
                addSpanID = self.multiline_apparatus_el.get('id')
                anchorID = f"span{self.multiline_apparatus_num}-close"
                srcs = ' '.join(srcs)
                self.multiline_apparatus_el.set('n', srcs)
                self.multiline_apparatus_el.set('to', anchorID)
                closeel = etree.XML(f'<anchor id="{anchorID}"  corresp="{addSpanID}" '
                                    f'rend="apparatus" ></anchor>')
                self.current_el.addnext(closeel)
                pfollowing = etree.XML('<p></p>')
                closeel.addnext(pfollowing)
                self.current_el = pfollowing
                self.multiline_apparatus_el = None
                self.in_multiline_apparatus = False  # though we are no longer in the mla
                self.iterate_runs(p, 2)  # Process remaining runs in this paragraph

        return paragraph_processed  # if this is false, the paragraph is processed as normal above

        # TODO: Deal with multiline apparatus

        # Blocked from getting here
        orig_current_el = self.current_el
        app_el = etree.XML('<p><app><rdg></rdg></app></p>')
        self.current_el = app_el.find('app').find('rdg')
        for ap in app_ps:
            if ap.text[0] == '{' or ap.text[0] == '}':
                continue
            self.iterate_runs(ap)
            if ap not in app_ps[-2:]:  # Not either the last line of app or the closing }
                lb = etree.XML('<lb/>')
                self.current_el.append(lb)
                self.current_el = lb
        for rn in app_ps[-1].runs:
            char_style = rn.style.name.lower()
            if "footnote" in char_style or "endnote" in char_style:
                note = self.endnotes.pop(0) if "endnote" in char_style.lower() else self.footnotes.pop(0)
                notedata = TextConverter.process_critical_note(note)
                sigla = self.edsig
                pgs = ''
                if notedata['lem']:
                    sigla = notedata['lem']['edsigs']
                    pgs = notedata['lem']['edpgs']
                if self.current_el.tag == 'lb':
                    self.current_el.getparent().set('wit', sigla)
                    if len(pgs) > 0:
                        self.current_el.getparent().set('n', pgs)
                else:
                    self.current_el.set('wit', sigla)
                    if len(pgs) > 0:
                        self.current_el.set('n', pgs)

                for vrnt in notedata['variants']:
                    rdg = '<rdg wit={} n={}>{}</rdg>'.format(vrnt['edsigs'], vrnt['edpgs'], vrnt['txt'])
                    rdg = etree.XML(rdg)
                    app_el.append(rdg)
        orig_current_el.addnext(app_el)
        self.current_el = app_el

    @staticmethod
    def createmilestone(char_style, mstxt):
        mstype = 'line' if 'line' in char_style.lower() else 'page'  # defaults to page
        msnum = mstxt.replace('[', '').replace(']', '')   # Default backup num if regex doesn't match
        mtch = re.match(r'\[?(Page|Line)\s+([^\]]+)\]?', mstxt, re.IGNORECASE)
        if mtch:
            mstype = mtch.group(1)
            msnum = mtch.group(2)
        else:  # In TCD some formatting weirdness in page milestones read as: [21-page Dg]
            mtch = re.match(r'\[?(\d)+\-page\s+([^\]]+)\]?', mstxt, re.IGNORECASE)
            if mtch:
                mstype = 'page'
                msnum = mtch.group(2) + '-' + mtch.group(1)
            # else:s
            #    logging.warning("No match for milestone parts in {}".format(mstxt))
        msel = getStyleElement(char_style)
        msel.set('unit', mstype)
        sep = '.' if '.' in msnum else '-'  # Do we need to check for more separators
        pts     = msnum.split(sep)
        if len(pts) > 1:
            msel.set('ed', pts[0])
            msel.set('n', pts[1])
        else:
            msel.set('n', pts[0])
        return msel

    def reset_current_el(self):
        """
        Resets the current element to the last element within the current div.
        This is used to get out of nested lists and verses

        :return:
        """
        if len(self.headstack) > 0:
            hdr = self.headstack[-1]
            children = hdr.getchildren()
            if len(children) == 0:
                self.current_el = hdr
            else:
                self.current_el = children[-1]

    def bodydivcheck(self):
        """
        This function ensures that the content of the body is wrapped in a chapter div for situations where
        there are not Header 1s in the document

        :return:
        """
        print("\rChecking body for chapter divs")
        bdivs = self.xmlroot.xpath('/*//text/body/div[@n="1"]')
        if len(bdivs) == 0:
            bd = self.xmlroot.xpath('/*//text/body')[0]
            bdchildren = bd.getchildren()
            div = etree.XML('<div n="1" id="b1"><head><num>2.1.</num> གཞུང་།</head></div>')
            bd.append(div)
            for bdchild in bdchildren:
                if bdchild.tag != 'head':
                    div.append(bdchild)

    def assignids(self):
        print("\rAssigning IDs");
        divs = self.xmlroot.xpath('/*//text//div')
        for divel in divs:
            headtxt = divel.xpath('./head[1]/num/text()')
            if headtxt:
                headtxt = headtxt[0]
                headlist = headtxt.split('.')
                mainsect = TextConverter.section_trans(headlist.pop(0))  # convert first number to letter
                headlist[0] = mainsect + headlist[0]
                myid = '-'.join(headlist)
                myid = myid.rstrip('-')  # remove any trailing dash
                divel.set('id', myid)
            else:
                head = divel.xpath('./head')
                head = etree.tostring(head[0]).decode('utf-8') if head else '???'
                head = "No num element in " + replace_entities(head)
                self.mylog(head)

        # Original code idea was to get the actual ancestor positions and use those, but his was problematic
        # for divel in divs:
        #     ancids = list(self.getSiblingPos(divel))
        #     textpt = ''  # To store the a/b/c for front, body or back
        #     for anc in divel.iterancestors():  # Iterates ancestors in reverse from current element up to TEI.2
        #         ancid = self.getSiblingPos(anc)  # returns number for div or a, b, c for front body or back
        #         if ancid.isnumeric():
        #             ancids.append(ancid)
        #         else:
        #             txtpt = ancid
        #             break   # break after front body or back to not count text or TEI.2
        #     ancids.reverse()
        #     if self.chapnum:
        #         ancids[0] = self.chapnum
        #     myid = txtpt + '-'.join(ancids)
        #     divel.set('id', myid)

    @staticmethod
    def getSiblingPos(el):
        atag = el.tag
        if atag == "front":
            return "a"
        if atag == "body":
            return "b"
        if atag == "back":
            return "c"
        anum = 1
        for sib in el.itersiblings('div', preceding=True):
            anum += 1
        return str(anum)

    @staticmethod
    def section_trans(fbbid):
        """
        Translates a front, body, or back number (1,2,3) to the appropriate letter (a, b, c) or vice versa
        :param fbbid:
        :return:
        """
        trans = ['', 'a', 'b', 'c']
        fbbid = int(fbbid) if isinstance(fbbid, str) and fbbid.isnumeric() else fbbid
        if isinstance(fbbid, str) and fbbid in trans:
            return trans.index(id)
        if isinstance(fbbid, int) and 0 < fbbid < 4:
            return trans[fbbid]
        return fbbid

    def tidyxml(self):
        empty_resp = self.xmlroot.xpath("//publicationStmt/respStmt/name[@n='agent' and not(text())]/parent::*")
        for resp in empty_resp:
            resp.getparent().remove(resp)

    def writexml(self):
        # Determine Name for Resulting XML file
        fname = self.current_file.replace('.docx', '.xml')
        fpth = os.path.join(self.outdir, fname)
        while os.path.isfile(fpth) and not self.overwrite:
            userin = input("The file {} already exists. Overwrite it (y/n/q): ".format(fname))
            if userin == 'y':
                break
            elif userin == 'n':
                fname = input("Enter a new file name: ")
                fpth = os.path.join(self.outdir, fname)
            else:
                exit(0)

        # Write XML File
        with open(fpth, "wb") as outfile:
            genid = self.textid.split('-text')[0] if '-text' in self.textid else self.textid
            bibid = genid + '-bib'
            # remove text document sub number for e.g. lccw-0353-1.docx
            # genid = re.sub(r'(-\d{4})-\d+', r'\1', genid)
            # Calculate bibl folder (first number of text id number)
            mtch = re.search(r'-(\d{4})', genid)
            fldr = mtch.group(1)[0] if mtch else '0'
            biblent = f"<!ENTITY {bibid} " \
                      f"SYSTEM \"../../{fldr}/{bibid}.xml\">" if self.args.bibl_entity is True else ""
            doc_type = f"<!DOCTYPE TEI.2 SYSTEM \"{self.dtdpath}xtib3.dtd\" [ \n" \
                f"\t<!ENTITY % thlnotent SYSTEM \"{self.dtdpath}catalog-refs.dtd\" > \n" \
                "\t%thlnotent;\n" \
                f"\t{biblent}\n]>"

            # Replace profile desc with entity
            pdentity = etree.Entity('thdlprofiledesc')
            pdesc = self.xmlroot.xpath('//profileDesc')[0]
            pdesc.addprevious(pdentity)
            pdesc.getparent().remove(pdesc)

            # Add tibbibl entity if there is a text id
            if self.textid and genid:
                tibsrc = etree.XML('<sourceDesc n="tibbibl"></sourceDesc>')
                if self.args.bibl_entity:
                    tibbibl_ent = etree.Entity(bibid)
                    tibsrc.append(tibbibl_ent)
                    tibsrc.tail = "\n"
                    docsrc = self.xmlroot.xpath('//sourceDesc')[0]
                    docsrc.addprevious(tibsrc)
            xmlstring = etree.tostring(self.xmlroot,
                                       pretty_print=True,
                                       encoding='utf-8',
                                       xml_declaration=True,
                                       doctype=doc_type)
            outfile.write(xmlstring)

    #  HELPER METHODS
    def get_previous_p(self, as_style=False):
        pind = self.pindex - 1
        if pind > -1:
            prevp = self.worddoc.paragraphs[pind]
            if as_style:
                return prevp.style.name
            else:
                return prevp
        return False

    # STATIC HELPER METHODS
    @staticmethod
    def mylog(msg):
        print(msg)
        logging.warning(msg)

    @staticmethod
    def getnumber(stynm):
        mtch = re.match(r'[^\d]+\s+(\d+)', stynm)
        if mtch:
            return mtch.group(1)
        return False

    @staticmethod
    def is_reg_p(style_name):
        # TODO: Check if I need to account for Paragraph Citation. Is that a "regular paragragh"?
        if "Paragraph" not in style_name and "Outline" not in style_name and "Normal" not in style_name:
            return False
        return True

    @staticmethod
    def getTag(element):
        return "%s:%s" % (element.prefix, re.match("{.*}(.*)", element.tag).group(1))

    def getTextFromCell(self, element):
        rels = self.worddoc.part.rels
        celltxt = ''
        for cellp in element.paragraphs:
            for cld in cellp._p:
                if cld.text:
                    celltxt += cld.text
                else:
                    # Get text from hyperlink
                    ctag = self.getTag(cld)
                    if ctag == "w:hyperlink":
                        rid = cld.get('{%s}id' % self.nsmap['r'])
                        for subc in cld:
                            ctag = self.getTag(subc)
                            if ctag == "w:r":
                                celltxt += subc.text
                        # Add on the href from the link in parenthese)
                        if rid in rels.keys():
                            href = rels[rid]._target
                            celltxt += ' (%s)' % href
        return celltxt

    def fn_is_annotation(self, fno):
        """
        Determines if a certain footnote is an annotation
        :param fno:
        :return:
        """
        is_annotation = fno['prev_el'] is not None and fno['prev_el'].text and fno['prev_el'].text[-1] == '}'
        if not is_annotation and fno['text']:
            is_annotation = re.search(ANNOTATION_PATTERN, fno['text'])
        if not is_annotation and fno['markup']:
            is_annotation = re.search(ANNOTATION_PATTERN, fno['markup'])
        return is_annotation

    def get_run_before_note(self, nnum):
        """
        Gets the immediately preceding text to a footnote
        :param nnum:
        :return:
        """
        prevrun = None
        for paragraph in self.worddoc.paragraphs:
            for run in paragraph.runs:
                if run.style.name == 'footnote reference':
                    fnnum, note = self.getFootnoteFromRefRun(run, False)
                    if fnnum == nnum:
                        if prevrun and prevrun.text:
                            # print(f"text before note {nnum}: {prevrun.text}")
                            return prevrun.text
                prevrun = run
        return prevrun

    def getFootnoteFromRefRun(self, run, include_note=True):
        # Get the footnote number from the run containing the "footnote reference"
        # Second element in footnote ref container has the number/id: <w:footnoteReference w:id="2"/>
        # The full attribute is {http://schemas.openxmlformats.org/wordprocessingml/2006/main}id
        # Easier just to pop the first key from the attribute dictionary of that element
        runel = run.element
        try:
            fnref = runel[1]
        except IndexError as ie:
            self.mylog("index error: {}".format(ie))
        fnnum = ""
        note = None

        if len(fnref.keys()) > 0:
            idkey = '{' + self.nsmap['w'] + '}id'
            fnnum = fnref.get(idkey)
            # print("doing footnote {}".format(fnnum))
            if include_note:
                note = self.footnotes[fnnum]

        return fnnum, note


class ConversionException(Exception):
    pass

