#!env/bin/python
"""
A convert that just addes digital pages to word docs
"""
from .baseconverter import BaseConverter
from os import path
import docx
from docx.text.run import Run
from docx.oxml.text.run import CT_R
import re


class DigitalPages(BaseConverter):
    def __init__(self, args, **kwargs):
        super().__init__(args, **kwargs)
        self.outfile = ""
        self.lines_per_page = 15
        self.tsek_per_line = 20
        self.digstyles = {
            'page': 'page number',  # 'page number',  # This is the digital page style name (page number repurposed)
            'line': 'line number'  # 'line number' # this is the digital line style name
        }
        # style_id_list = [digpage_style_id, digline_style_id]
        self.pgct = 1  # Page start 1 (to insert first page and line markers)
        self.lnct = 0  # Line start 0 (increment after adding so we can use = below)
        self.ct = 0
        self.tmplts = {
            'page': 'tdp',
            'line': 'tdl'
        }
        self.tsekpattern = r'[\u0F00-\u0F14\u0F3A-\u0F3D\u0FD2-\u0FD8\s]+'  # at least one of the Tibetan punctuation or space-like characters
        self.tskcount = 0
        self.donefirst = False

    def convertdoc(self):
        self.current_file_path = path.join(self.indir, self.current_file)
        self.worddoc = docx.Document(self.current_file_path)
        self.merge_runs()
        totalpct = len(self.worddoc.paragraphs)
        pct = 0
        # Reset class variables for each document
        self.pgct = 1
        self.lnct = 0
        self.ct = 0
        self.tskcount = 0
        self.donefirst = False

        print("Inserting placeholders ...")

        foundhead = False
        for p in self.worddoc.paragraphs:
            pct += 1
            print("\rDoing paragraph {} of {}        ".format(pct, totalpct), end="")
            psnm = p.style.name
            # Don't start counting until we get the first header (usually front or body)
            if not foundhead and 'Heading' not in psnm:
                continue
            foundhead = True
            # After that ignore headers and only proc ess non-headers (skip any of the following styles)
            doinsert = True
            for igsn in ['Heading', 'Page Number', 'Line Number']:
                if igsn in psnm or igsn.lower() in psnm:
                    doinsert = False
                    break
            if doinsert:
                self.insert_ms(p)
                self.apply_styles(p)
        print("\n")
        self.outfile = path.join(self.outdir, self.current_file.replace('.doc', '-out.doc'))
        self.worddoc.save(self.outfile)

    def insert_ms(self, p):
        for r in p.runs:
            inserts = []
            rtxt = r.text
            endindex = -1
            for m in re.finditer(self.tsekpattern, rtxt):
                endindex = m.end()
                if m.start() == 0:
                    continue
                self.tskcount += 1
                if self.tskcount == self.tsek_per_line:
                    self.lnct += 1
                    self.tskcount = 0
                    if self.lnct == self.lines_per_page:
                        self.pgct += 1
                        self.lnct = 0
                        inserts.append(('page', f"{self.pgct}", m))
                        inserts.append(('line', f"{self.pgct}.{self.lnct + 1}", m))
                    else:
                        inserts.append(('line', f"{self.pgct}.{self.lnct + 1}", m))

            inserts.reverse()
            for ins in inserts:
                mstype, msnum, m = ins
                ms = f"[{self.tmplts[mstype]} {msnum}]"
                rtxt = rtxt[:m.end()] + ms + rtxt[m.end():]

            r.text = rtxt
            if not self.donefirst:
                rtxt = f"[{self.tmplts['page']} 1][{self.tmplts['line']} 1.1]" + rtxt
                r.text = rtxt
                self.donefirst = True

    def apply_styles(self, p):
        newp = p.insert_paragraph_before('', p.style)
        for r in p.runs:
            rtxt = r.text
            pgrex = '\[' + self.tmplts["page"] + '\s+\d+\]'
            mtchs = [m for m in re.finditer(pgrex, rtxt)]
            if len(mtchs) > 0:
                newp.add_run(rtxt[:mtchs[0].start()], r.style)
                for mi, mtch in enumerate(mtchs):
                    ms = rtxt[mtch.start():mtch.end()]
                    ms = ms.replace(self.tmplts["page"] + ' ', '')
                    newp.add_run(ms, self.worddoc.styles[self.digstyles['page']])
                    endindex = mtchs[mi + 1].start() if mi < len(mtchs) - 1 else len(rtxt)
                    postrun = rtxt[mtch.end():endindex]
                    newp.add_run(postrun, r.style)
            else:
                newp.add_run(r.text, r.style)
        self.delete_paragraph(p)
        # Do Line milestones
        newp2 = newp.insert_paragraph_before('', newp.style)
        for r in newp.runs:
            rtxt = r.text
            pgrex = '\[' + self.tmplts["line"] + '\s+\d+\.\d+\]'
            mtchs = [m for m in re.finditer(pgrex, rtxt)]
            if len(mtchs) > 0:
                # Add text before first match
                newp2.add_run(rtxt[:mtchs[0].start()], r.style)
                for mi, mtch in enumerate(mtchs):
                    # add milestone run
                    ms = rtxt[mtch.start():mtch.end()]
                    ms = ms.replace(self.tmplts["line"] + ' ', '')
                    newp2.add_run(ms, self.worddoc.styles[self.digstyles['line']])
                    # Add post milestone run
                    endindex = mtchs[mi + 1].start() if mi < len(mtchs) - 1 else len(rtxt)
                    postrun = rtxt[mtch.end():endindex]
                    newp2.add_run(postrun, r.style)
            else:
                newp2.add_run(r.text, r.style)
        self.delete_paragraph(newp)

    def add_run(self, r, txt, style_name):
        runel = CT_R()
        r._element.addnext(runel)
        newrun = Run(runel, r._parent)
        newrun.text = txt
        newrun.style = self.worddoc.styles[style_name]
        return newrun

    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
