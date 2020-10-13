#!env/bin/python
"""
The base object inherited by all converters
"""
import os
import logging
import re
import zipfile
import docx
from lxml import etree
from datetime import date

TEMPLATE_FOLDER = 'templates'


class BaseConverter:

    def __init__(self, args):
        self.files = []
        self.current_file = ''
        self.current_file_path = ''
        self.indir = args.indir
        if not os.path.isdir(self.indir):
            raise NotADirectoryError("The in path, {}, is not a directory".format(self.indir))
        self.getfiles()
        self.outdir = args.out
        if not os.path.isdir(self.outdir):
            raise NotADirectoryError("The out path, {}, is not a directory".format(self.outdir))
        self.metafields = args.metafields if args.metafields else False
        self.template = args.template
        self.xmltemplate = ''
        self.dtdpath = args.dtdpath
        self.debug = args.debug if args.debug else False
        self.worddoc = None
        self.metatable = None
        self.footnotes = []
        self.endnotes = []
        self.xmlroot = None
        self.headstack = []
        self.current_el = None

        self.log = args.log
        self.loglevel = logging.DEBUG if self.debug else logging.WARN
        logging.basicConfig(level=self.loglevel)

    def getfiles(self):
        for sfile in os.listdir(self.indir):
            if sfile.endswith(".docx") and not sfile.startswith('~'):
                self.files.append(sfile)

    def setlog(self):
        # fname = os.path.split(fname)[1].replace('.docx', '') + '.log'
        logpath = os.path.join(self.log, self.current_file.replace('docx', 'log'))
        print(logpath)
        if self.debug:
            print("Log file for {} is: {}".format(self.current_file, logpath))
        loghandler = logging.FileHandler(logpath, 'w')
        log = logging.getLogger()
        for hdlr in log.handlers[:]:
            log.removeHandler(hdlr)
        log.addHandler(loghandler)

    def convert(self):
        for fl in self.files:
            print("doing file: {}".format(fl))
            self.current_file = fl
            self.setlog()
            self.convertdoc()
            self.writexml()

    def convertdoc(self):
        self.current_file_path = os.path.join(self.indir, self.current_file)
        self.worddoc = docx.Document(self.current_file_path)
        print("it has {} paragraphs".format(len(self.worddoc.paragraphs)))
        self.processFootEndNotes()
        self.createxml()
        # TODO: add more conversion code here. Iterate through paras then runs. Using stack method
        for p in self.worddoc.paragraphs:
            if isinstance(p, docx.text.paragraph.Paragraph):
                self.convertpara(p)
            else:
                self.mywarning("Warning: paragraph ({}) is not a docx paragraph cannot convert".format(p))

    def processFootEndNotes(self):
        zipdoc = zipfile.ZipFile(self.current_file_path)

        # write content of endnotes.xml into self.footnotes[]
        fnotestxt = zipdoc.read('word/footnotes.xml')
        xml_fn_root = etree.fromstring(fnotestxt)
        fnindex = 0
        fnotes = xml_fn_root.findall('w:footnote', xml_fn_root.nsmap)
        for f in fnotes:
            if fnindex > 1:
                text = f.findall('.//w:t', xml_fn_root.nsmap)
                s = ""
                for t in text:
                    s += t.text
                self.footnotes.append(s)
            fnindex += 1

        # write content of endnotes.xml into self.endnotes[]
        xml_content = zipdoc.read('word/endnotes.xml')
        xml_en_root = etree.fromstring(xml_content)
        enindex = 0
        enotes = xml_en_root.findall('w:endnote', xml_en_root.nsmap)
        for f in enotes:
            if enindex > 1:
                text = f.findall('.//w:t', xml_en_root.nsmap)
                s = ""
                for t in text:
                    s += t.text
                self.endnotes.append(s)
            enindex += 1
        zipdoc.close()

    def createxml(self):
        with open(os.path.join(TEMPLATE_FOLDER, self.template), 'r') as tempstream:
            self.xmltemplate = tempstream.read()
            self.metatable = self.worddoc.tables[0] if len(self.worddoc.tables) else False
            if self.metatable:
                self.createmeta()  # Separated out to be overriden. Creates the XML string with metadata inserted
            # Take self.xmltemplate the xml string and convert to an XML object
            xmldoc = bytes(bytearray(self.xmltemplate, encoding='utf-8'))
            # create lxml element tree from metadata info
            parser = etree.XMLParser(ns_clean=True, recover=True, encoding='utf-8')
            self.xmlroot = etree.fromstring(xmldoc, parser)

    def createmeta(self):
        """
        This table creates the xmlstring by filling in the template string with the table values based on the
        table labels. It uses the classes self.metatable property which is the Word docs metatable and
        self.template which is the string read in from the XML template file indicated in the settings

        NOTE: This function should be overridden by particular template converters to customize
        TODO: Create a generalized version to put here and use this for the chris_old template
        :return:
        """
        wordtable = self.metatable
        # Fill out metadata matching on string in wordtable with {strings} in template (teiHeader.dat)
        xmltext = self.xmltemplate.replace("{Digital Creation Date}", str(date.today()))
        problems_on = False
        tablerows = len(wordtable.rows)
        for rwnum in range(0, tablerows):
            try:
                if wordtable._column_count == 2:
                    label = wordtable.cell(rwnum, 0).text.strip()
                    rowval = wordtable.cell(rwnum, 1).text.strip()
                elif wordtable._column_count == 3:
                    label = wordtable.cell(rwnum, 0).text.strip()
                    rowval = wordtable.cell(rwnum, 1).text.strip()
                    msg = "***** NOTICE: Need to update code for processing 3 column metadata tables!!!!! *******"
                    # print(msg)
                    logging.info(msg)
                else:
                    raise (ValueError('Metadata table does not have the right number of columns. Must be 2 or 3'))

                # All Uppercase are Headers in the table skip
                if label.isupper():
                    if label == 'PROBLEMS':
                        problems_on = True
                    continue  # Skip labels
                if problems_on:
                    if '<encodingDesc>' not in xmltext:
                        xmltext = xmltext.replace('</fileDesc>',
                                                    '</fileDesc><encodingDesc><editorialDecl ' +
                                                    'n="problems"><interpretation n="{}">'.format(label) +
                                                    '<p>{}</p>'.format(rowval) +
                                                    '</interpretation></editorialDecl></encodingDesc>')
                    else:
                        xmltext = xmltext.replace('</interpretation></editorialDecl>',
                                                    '</interpretation><interpretation n="{}">'.format(label) +
                                                    '<p>{}</p>'.format(rowval) +
                                                    '</interpretation></editorialDecl>')
                temppt = label.split(' (')  # some rows have " (if applicable)" or possible some other instruction
                label = temppt[0].replace('\xa0', ' ').strip()
                label = label.replace(' (if applicable)', '')
                label = label.replace('Call་number', 'Call-number')
                srclbl = "{" + label + "}"
                xmltext = xmltext.replace(srclbl, rowval)

            except IndexError as e:
                logging.error("Index error in iterating wordtable: {}".format(e))
            except TypeError as e:
                logging.error("Type error in iterating wordtable: {}".format(e))

        self.xmltemplate = re.sub(r'{([^}]+)}', r'<!--\1-->', xmltext)

    def convertpara(self, p):
        style_name = p.style.name
        headmtch = re.match(r'^Heading (?:Tibetan\s*)?(\d+)[\,\s]*(Front|Body|Back)?', style_name)
        if headmtch:
            self.doheader(p, headmtch)
        else:
            # TODO: fill in for regular paragraph processing
            pass

    def doheader(self, p, headmtch):
        """
        Method to convert heading styles into structural elements either front, body, bock or divs
        Called from convertpara() method above

        :param p: the word docx paragraph object
        :param headmtch: the re.match object group(1) is heading level, group(2) is front, body, back if it is one of those
        :return: none
        """
        hlevel = int(headmtch.group(1))
        # TODO: need to parse the p element below in case there is internal markup to put in head
        #  (i.e. don't just use p.text but it may have children)
        htext = p.text
        mtch = re.match(r'^((\d+\.?)+)', htext)
        if mtch:
            hnum = mtch.group(1)
            htext = htext.replace(hnum, '<num>{}</num>'.format(hnum))
        style_name = p.style.name
        if hlevel == 0:
            # If level is 0, its front body or back, create element and clear head stack
            fbbel = etree.XML('<{0}><head>{1}</head></{0}>'.format(headmtch.group(2).lower(), htext))
            self.xmlroot.find('text').append(fbbel)
            self.current_el = fbbel
            self.headstack = [fbbel]
        else:
            # Otherwise we are already in front, body, or back, so create div
            currlevel = len(self.headstack) - 1  # subtract one bec. div 0 is at top of stack
            # TODO: need to parse the p element in case there is internal markup to put in head
            hdiv = etree.XML('<div n="{}"><head>{}</head><p></p></div>'.format(hlevel, htext))
            # if it's the next level deeper
            if hlevel > currlevel:
                # if new level is higher than the current level just add it to current
                if hlevel - currlevel > 1:
                    self.mywarning("Warning: Heading level skipped for {}".format(style_name, htext))
                self.current_el.append(hdiv)
                self.current_el = hdiv
                self.headstack.append(hdiv)
            # if it's the same level as current
            elif hlevel == currlevel:
                self.current_el.getparent().append(hdiv)
                self.current_el = hdiv
                self.headstack[-1] = hdiv
            # Otherwise it's a higher level
            else:
                # because front, body, back is 0th element use hlevel to splice array
                self.headstack = self.headstack[0:hlevel]
                self.headstack[-1].append(hdiv)
                self.headstack.append(hdiv)
                self.current_el = hdiv

    def writexml(self):
        # Determine Name for Resulting XML file
        fname = self.current_file.replace('.docx', '.xml')
        fpth = os.path.join(self.outdir, fname)
        while os.path.isfile(fpth):
            userin = input("The file {} already exists. Overwrite it (y/n/q): ".format(fname))
            if userin == 'y':
                break
            elif userin == 'n':
                fname = input("Enter a new file name: ")
                fpth = os.path.join(self.outdir, fname)
            else:
                exit(0)

        # Write XML File
        with open(fpth, "wb") as outfile:
            docType = "<!DOCTYPE TEI.2 SYSTEM \"{}xtib3.dtd\">".format(self.dtdpath)
            xmlstring = etree.tostring(self.xmlroot,
                                        pretty_print=True,
                                        encoding='utf-8',
                                        xml_declaration=True,
                                        doctype=docType)
            outfile.write(xmlstring)

    @staticmethod
    def mywarning(msg):
        print(msg)
        logging.warning(msg)